"""
Certificate document generation module using python-docx
Generates certificates from Word templates with placeholder replacement and signature support.
"""

import os
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

from django.conf import settings
from django.core.files.base import ContentFile
from docx import Document
from docx.shared import Inches, Pt

from .models import GeneratedCertificate


class CertificateDocumentGenerator:
    """
    Generate certificate Word documents from templates with dynamic content and optional signatures.
    Uses python-docx to manipulate .docx templates.
    """
    
    def __init__(self, certificate: GeneratedCertificate):
        self.certificate = certificate
        self.template = certificate.template
        self.data = certificate.certificate_data.copy()
        self.signature = certificate.signature_official if certificate.include_signature else None
    
    def _prepare_data(self):
        """Add computed fields and missing data before replacement"""
        # Add recipient_name from certificate model
        self.data['recipient_name'] = self.certificate.recipient_name
        
        # Add current date fields
        now = datetime.now()
        self.data['day'] = now.strftime('%d')
        self.data['month'] = now.strftime('%B')
        self.data['year'] = now.strftime('%Y')
    
    def generate_from_template(self, template_path):
        """
        Generate certificate by loading template and replacing placeholders.
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        # Prepare data
        self._prepare_data()
        
        # Open template
        doc = Document(template_path)
        
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            # Get combined text from all runs
            combined_text = ''.join(run.text for run in paragraph.runs)
            new_text = combined_text
            
            # Replace all placeholders
            for key, value in self.data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value) if value else '')
            
            # Update if changed
            if new_text != combined_text:
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ''
                # Add new text
                paragraph.add_run(new_text)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        combined_text = ''.join(run.text for run in paragraph.runs)
                        new_text = combined_text
                        for key, value in self.data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, str(value) if value else '')
                        if new_text != combined_text:
                            for run in paragraph.runs:
                                run.text = ''
                            paragraph.add_run(new_text)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def save_generated_document(self, template_path, certificate_obj=None):
        """Generate document and save to certificate."""
        if certificate_obj is None:
            certificate_obj = self.certificate
        
        try:
            doc_buffer = self.generate_from_template(template_path)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            recipient_name = self.certificate.recipient_name.replace(' ', '_')[:30]
            template_name = self.template.template_type
            docx_filename = f"{template_name}_{recipient_name}_{timestamp}.docx"
            
            certificate_obj.docx_file.save(docx_filename, ContentFile(doc_buffer.getvalue()), save=True)
            return certificate_obj.docx_file.path
            
        except Exception as e:
            raise e
    
    def save_files(self, certificate_obj=None):
        """Save certificate using template file. Skips if no template file configured."""
        if certificate_obj is None:
            certificate_obj = self.certificate
        
        if not self.template.template_file:
            # No template file configured - skip file generation
            # The HTML template can be used for certificate generation
            return None
        
        template_path = self.template.template_file.path
        return self.save_generated_document(template_path, certificate_obj)