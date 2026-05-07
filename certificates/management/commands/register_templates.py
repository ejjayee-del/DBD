import os
import re
from pathlib import Path
from docx import Document
from django.core.management.base import BaseCommand
from django.core.files import File
from certificates.models import CertificateTemplate


class Command(BaseCommand):
    help = 'Register certificate templates from template files'

    def add_arguments(self, parser):
        parser.add_argument(
            '--template-dir',
            type=str,
            default='media/templates',
            help='Path to template directory'
        )
        parser.add_argument(
            '--replace',
            action='store_true',
            help='Replace existing templates'
        )

    def handle(self, *args, **options):
        template_dir = Path(options['template_dir'])
        replace = options['replace']

        if not template_dir.exists():
            self.stdout.write(
                self.style.ERROR(f'Template directory {template_dir} does not exist')
            )
            return

        # Find all .docx files
        docx_files = list(template_dir.glob('*.docx'))

        if not docx_files:
            self.stdout.write(
                self.style.WARNING('No .docx files found in template directory')
            )
            return

        self.stdout.write(
            self.style.SUCCESS(f'Found {len(docx_files)} template file(s)')
        )

        for filepath in docx_files:
            template_name = filepath.stem  # Filename without extension

            # Extract placeholders from document
            try:
                doc = Document(filepath)
                placeholders = self._extract_placeholders(doc)
                self.stdout.write(f'\nProcessing: {template_name}')
                self.stdout.write(f'  Placeholders found: {", ".join(placeholders)}')

                # Check if template already exists
                template = CertificateTemplate.objects.filter(
                    template_type=template_name.lower().replace(' ', '_')
                ).first()

                if template and not replace:
                    self.stdout.write(
                        self.style.WARNING(
                            f'  Template "{template_name}" already exists (use --replace to override)'
                        )
                    )
                    continue

                # Remove old template if replacing
                if template and replace:
                    template.delete()
                    self.stdout.write(f'  Deleted existing template: {template_name}')

                # Create new template
                template_type = template_name.lower().replace(' ', '_')
                # Map to valid template types
                type_mapping = {
                    'barangay_clearance': 'barangay_clearance',
                    'certificate_of_residency': 'residency',
                    'certificate_of_indigency': 'indigency',
                    'business_permit': 'business_permit',
                    'first_time_job_seeker': 'first_time_job_seeker',
                    'solo_parent_certificate': 'solo_parent',
                    'senior_citizen_certificate': 'senior_citizen',
                    'pwd_certificate': 'pwd',
                }
                template_type = type_mapping.get(template_type, 'other')

                template = CertificateTemplate.objects.create(
                    template_type=template_type,
                    template_name=template_name,
                    description=f'Auto-registered template from {filepath.name}',
                    is_active=True,
                )

                # Save the template file
                with open(filepath, 'rb') as f:
                    template.template_file.save(filepath.name, File(f), save=True)

                # Create fields from placeholders
                required_fields = {}
                for placeholder in placeholders:
                    # Try to infer field type from name
                    field_type = self._infer_field_type(placeholder)
                    
                    required_fields[placeholder] = {
                        'label': placeholder.replace('_', ' ').title(),
                        'type': field_type,
                        'required': True,
                    }

                template.required_fields = required_fields
                template.save()

                self.stdout.write(
                    self.style.SUCCESS(f'✓ Registered template: {template_name}')
                )

            except Exception as e:
                self.stdout.write(
                    self.style.ERROR(f'  Error processing {template_name}: {str(e)}')
                )

        self.stdout.write(self.style.SUCCESS('\nTemplate registration complete!'))

    def _extract_placeholders(self, doc):
        """Extract placeholder names from document"""
        placeholders = set()
        
        # Pattern for {{placeholder}} style markers
        pattern = r'\{\{([a-zA-Z0-9_]+)\}\}'

        # Check paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            placeholders.update(matches)

        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        placeholders.update(matches)

        return sorted(list(placeholders))

    def _infer_field_type(self, field_name):
        """Infer field type from field name"""
        field_lower = field_name.lower()

        if 'email' in field_lower:
            return 'email'
        elif 'phone' in field_lower or 'contact' in field_lower or 'mobile' in field_lower:
            return 'phone'
        elif 'date' in field_lower or 'born' in field_lower:
            return 'date'
        elif 'address' in field_lower or 'location' in field_lower:
            return 'textarea'
        else:
            return 'text'
