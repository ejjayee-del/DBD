import re
from pathlib import Path

from django import forms
from docx import Document

from .models import CertificateDocument, CertificateTemplate, GeneratedCertificate


class CertificateGenerationForm(forms.ModelForm):
    """Dynamic form for generating certificates based on template"""
    
    class Meta:
        model = GeneratedCertificate
        fields = ['recipient_name', 'recipient_email', 'recipient_contact', 'include_signature', 'signature_official']
        widgets = {
            'recipient_name': forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Full Name'}),
            'recipient_email': forms.EmailInput(attrs={'class': 'form-control', 'placeholder': 'Email (Optional)'}),
            'recipient_contact': forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Contact Number (Optional)'}),
            'include_signature': forms.CheckboxInput(attrs={'class': 'form-check-input'}),
            'signature_official': forms.Select(attrs={'class': 'form-control'}),
        }


class DynamicCertificateForm(forms.Form):
    """Dynamically generated form based on certificate template fields"""
    
    PLACEHOLDER_PATTERN = re.compile(r'\{\{([a-zA-Z0-9_]+)\}\}')
    
    def _infer_field_type(self, field_name):
        """Infer field type from placeholder name"""
        field_lower = field_name.lower()
        if 'email' in field_lower:
            return 'email'
        elif 'phone' in field_lower or 'contact' in field_lower or 'mobile' in field_lower:
            return 'phone'
        elif 'date' in field_lower or 'born' in field_lower or 'issued' in field_lower:
            return 'date'
        elif 'address' in field_lower or 'location' in field_lower or 'purpose' in field_lower:
            return 'textarea'
        elif 'civil_status' in field_lower or 'civil status' in field_lower:
            return 'select'
        else:
            return 'text'
    
    def _extract_docx_placeholders(self, template):
        """Extract {{placeholder}} patterns from .docx template file"""
        if not template.template_file:
            return {}
        
        filepath = template.template_file.path
        if not Path(filepath).exists():
            return {}
        
        try:
            doc = Document(filepath)
            placeholders = set()
            
            # Check paragraphs
            for paragraph in doc.paragraphs:
                matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
                placeholders.update(matches)
            
            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
                            placeholders.update(matches)
            
            # Build field config from placeholders
            fields = {}
            for placeholder in sorted(placeholders):
                # Skip computed placeholders added at generation time
                if placeholder in ('day', 'month', 'year'):
                    continue
                
                field_config = {
                    'label': placeholder.replace('_', ' ').title(),
                    'type': self._infer_field_type(placeholder),
                    'required': True,
                }
                
                # Add civil status options for civil_status fields
                if placeholder.lower() in ('civil_status', 'civil status'):
                    field_config['choices'] = 'Single,Married,Widowed,Separated,Live-in,Common Law'
                
                fields[placeholder] = field_config
            return fields
        except Exception:
            return {}
    
    def __init__(self, template, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.template = template
        
        # Try to get fields from .docx template file first
        docx_fields = self._extract_docx_placeholders(template)
        
        if docx_fields:
            all_fields = docx_fields
        else:
            # Fall back to database fields
            all_fields = template.get_all_fields()
        
        for field_name, field_config in all_fields.items():
            field_type = field_config.get('type', 'text')
            field_label = field_config.get('label', field_name)
            is_required = field_config.get('required', False)
            help_text = field_config.get('help_text', '')
            placeholder = field_config.get('placeholder', '')
            
            if field_type == 'text':
                self.fields[field_name] = forms.CharField(
                    label=field_label,
                    required=is_required,
                    help_text=help_text,
                    widget=forms.TextInput(attrs={
                        'class': 'form-control',
                        'placeholder': placeholder or field_label
                    })
                )
            elif field_type == 'textarea':
                self.fields[field_name] = forms.CharField(
                    label=field_label,
                    required=is_required,
                    help_text=help_text,
                    widget=forms.Textarea(attrs={
                        'class': 'form-control',
                        'rows': 4,
                        'placeholder': placeholder or field_label
                    })
                )
            elif field_type == 'date':
                self.fields[field_name] = forms.DateField(
                    label=field_label,
                    required=is_required,
                    help_text=help_text,
                    widget=forms.DateInput(attrs={
                        'class': 'form-control',
                        'type': 'date'
                    })
                )
            elif field_type == 'email':
                self.fields[field_name] = forms.EmailField(
                    label=field_label,
                    required=is_required,
                    help_text=help_text,
                    widget=forms.EmailInput(attrs={
                        'class': 'form-control',
                        'placeholder': placeholder or field_label
                    })
                )
            elif field_type == 'phone':
                self.fields[field_name] = forms.CharField(
                    label=field_label,
                    required=is_required,
                    help_text=help_text,
                    widget=forms.TextInput(attrs={
                        'class': 'form-control',
                        'type': 'tel',
                        'placeholder': placeholder or field_label
                    })
                )
            elif field_type == 'number':
                self.fields[field_name] = forms.DecimalField(
                    label=field_label,
                    required=is_required,
                    help_text=help_text,
                    widget=forms.NumberInput(attrs={
                        'class': 'form-control',
                        'placeholder': placeholder or field_label
                    })
                )
            elif field_type == 'select':
                choices_str = field_config.get('choices', '')
                choices = [('', '-- Select --')] + [(c.strip(), c.strip()) for c in choices_str.split(',') if c.strip()]
                self.fields[field_name] = forms.ChoiceField(
                    label=field_label,
                    required=is_required,
                    choices=choices,
                    help_text=help_text,
                    widget=forms.Select(attrs={'class': 'form-control'})
                )
            elif field_type == 'radio':
                choices_str = field_config.get('choices', '')
                choices = [(c.strip(), c.strip()) for c in choices_str.split(',') if c.strip()]
                self.fields[field_name] = forms.ChoiceField(
                    label=field_label,
                    required=is_required,
                    choices=choices,
                    help_text=help_text,
                    widget=forms.RadioSelect(attrs={'class': 'form-check-input'})
                )
                
class CertificateDocumentForm(forms.ModelForm):
    class Meta:
        model = CertificateDocument
        fields = ['title', 'content']