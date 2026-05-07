# Standard library imports
from datetime import datetime, date
from io import BytesIO

# Third-party imports
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

# Django imports
from django import forms
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.http import FileResponse, HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.db.models import Q
from django.utils import timezone
from django.views.decorators.http import require_http_methods
from django_summernote.widgets import SummernoteWidget

# Local imports
from transactions.models import ActivityLog
from .models import CertificateDocument, CertificateTemplate, GeneratedCertificate, CertificateRequest
from .forms import CertificateGenerationForm, DynamicCertificateForm


class EditableCertificateForm(forms.ModelForm):
    """Form for editing certificate content in the system"""
    class Meta:
        model = GeneratedCertificate
        fields = ['certificate_data']
        widgets = {
            'certificate_data': SummernoteWidget(),
        }


@login_required
@require_http_methods(["GET"])
def certificate_types_view(request):
    """Display all available certificate types"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to generate certificates.')
        return redirect('dashboard')
    
    query = request.GET.get('q', '').strip()
    certificates = CertificateTemplate.objects.filter(is_active=True)

    if query:
        certificates = certificates.filter(
            Q(template_name__icontains=query) |
            Q(description__icontains=query) |
            Q(template_type__icontains=query)
        )

    certificates = list(certificates)
    kp_form_types = ['hearing_notice', 'complaint_form', 'settlement']
    
    # Define categories mapping
    category_mapping = {
        'barangay_clearance': 'Basic Residency Certificates',
        'residency': 'Basic Residency Certificates',
        'non_residency': 'Basic Residency Certificates',
        'indigency': 'Basic Residency Certificates',
        'Food_Assistance': 'Assistance Certificates',
        'Financial_Assistance': 'Assistance Certificates',
        'Program_Certificate': 'Program & Activity Certificates',
        'Participation': 'Program & Activity Certificates',
        'senior_citizen': 'Special Groups Certificates',
        'pwd': 'Special Groups Certificates',
        'solo_parent': 'Special Groups Certificates',
        'first_time_job_seeker': 'Special Groups Certificates',
        'death_certificate': 'Legal & Official Forms',
        'complaint_form': 'Legal & Official Forms',
        'hearing_notice': 'Legal & Official Forms',
        'settlement': 'Legal & Official Forms',
        'appearance': 'Legal & Official Forms',
        'duration': 'Legal & Official Forms',
        'business_permit': 'Business & Ownership Certificates',
        'animal_ownership': 'Business & Ownership Certificates',
        'certification': 'Health & Other Certificates',
        'other': 'Health & Other Certificates',
    }
    
    # Group certificates by category
    categorized_certificates = {}
    for cert in certificates:
        category = category_mapping.get(cert.template_type, 'Other Certificates')
        if category not in categorized_certificates:
            categorized_certificates[category] = []
        categorized_certificates[category].append(cert)
    
    context = {
        'certificates': certificates,
        'categorized_certificates': categorized_certificates,
        'kp_form_types': kp_form_types,
        'query': query,
    }
    
    return render(request, 'certificates/certificate_types.html', context)


@login_required
@require_http_methods(["GET", "POST"])
def generate_certificate_view(request, template_id):
    """Generate a new certificate from template"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to generate certificates.')
        return redirect('dashboard')
    
    template = get_object_or_404(CertificateTemplate, id=template_id, is_active=True)
    
    if request.method == 'POST':
        form = CertificateGenerationForm(request.POST)
        dynamic_form = DynamicCertificateForm(template, request.POST)
        
        if form.is_valid() and dynamic_form.is_valid():
            # Convert date objects to strings for JSON serialization
            certificate_data = {}
            for key, value in dynamic_form.cleaned_data.items():
                if isinstance(value, (date, datetime)):
                    certificate_data[key] = value.isoformat()
                else:
                    certificate_data[key] = value
            
            # Create certificate
            certificate = form.save(commit=False)
            certificate.template = template
            certificate.created_by = request.user
            certificate.status = 'generated'
            certificate.certificate_data = certificate_data
            certificate.save()
            
            # Generate PDF and Word files - lazy import
            try:
                from .document_generator import CertificateDocumentGenerator
                generator = CertificateDocumentGenerator(certificate)
                generator.save_files(certificate)
            except ImportError as e:
                # Log the exception and notify user
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Document generation failed: {e}")
                messages.warning(request, 'Certificate saved but document generation is currently unavailable. Please contact an administrator.')
            
            # Log the activity
            ActivityLog.objects.create(
                user=request.user,
                log_type='certificate_generated',
                action_description=f"Generated {template.template_name} for {certificate.recipient_name}",
                object_id=certificate.id,
            )
            
            messages.success(request, 'Certificate generated successfully!')
            return redirect('certificates:preview', certificate_id=certificate.id)
    else:
        form = CertificateGenerationForm()
        dynamic_form = DynamicCertificateForm(template)
    
    context = {
        'template': template,
        'form': form,
        'dynamic_form': dynamic_form,
    }
    
    return render(request, 'certificates/generate_certificate.html', context)


@login_required
@require_http_methods(["GET"])
def preview_certificate_view(request, certificate_id):
    """Preview generated certificate"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    # Check permission
    if certificate.created_by != request.user and not request.user.can_view_history:
        messages.error(request, 'You do not have permission to view this certificate.')
        return redirect('dashboard')
    
    # Render the HTML template with certificate data
    from django.template import Template, Context
    template = Template(certificate.template.html_template)
    context = Context(certificate.certificate_data)
    rendered_content = template.render(context)
    
    context = {
        'certificate': certificate,
        'rendered_content': rendered_content,
    }
    
    return render(request, 'certificates/preview_certificate.html', context)


@login_required
@require_http_methods(["GET"])
def download_certificate_view(request, certificate_id, file_type='docx'):
    """Download certificate as Word document"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    # Check permission
    if certificate.created_by != request.user and not request.user.can_view_history:
        messages.error(request, 'You do not have permission to view this certificate.')
        return redirect('dashboard')
    
    if certificate.docx_file:
        file_path = certificate.docx_file.path
        with open(file_path, 'rb') as f:
            return FileResponse(f, 
                              content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              as_attachment=True,
                              filename=f"{certificate.recipient_name}_{certificate.template.template_name}.docx")
    
    messages.error(request, 'Certificate file not found.')
    return redirect('certificates:preview', certificate_id=certificate_id)


@login_required
@require_http_methods(["GET"])
def print_certificate_view(request, certificate_id):
    """Mark certificate as printed and record it"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to print certificates.')
        return redirect('dashboard')
    
    # Update certificate status
    certificate.status = 'printed'
    certificate.printed_date = timezone.now()
    certificate.printed_by = request.user
    certificate.save()
    
    # Log the activity
    ActivityLog.objects.create(
        user=request.user,
        log_type='certificate_printed',
        action_description=f"Printed {certificate.template.template_name} for {certificate.recipient_name}",
        object_id=certificate.id,
    )
    
    messages.success(request, 'Certificate marked as printed.')
    return redirect('certificates:preview', certificate_id=certificate_id)


@login_required
@require_http_methods(["GET"])
def list_certificates_view(request):
    """List all generated certificates"""
    if not request.user.can_view_history:
        messages.error(request, 'You do not have permission to view certificates.')
        return redirect('dashboard')
    
    certificates = GeneratedCertificate.objects.select_related('template', 'created_by').all()
    
    # Filter by status if provided
    status = request.GET.get('status')
    if status:
        certificates = certificates.filter(status=status)
    
    # Pagination
    paginator = Paginator(certificates, 25)  # 25 per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'certificates': page_obj,
    }
    
    return render(request, 'certificates/list_certificates.html', context)


@staff_member_required
def download_document_as_docx(request, document_id):
    """Download certificate document as Word file"""
    document = get_object_or_404(CertificateDocument, id=document_id)
    
    # Create Word document
    doc = Document()
    
    # Add title
    doc.add_heading(document.title, 0)
    
    # Parse HTML content from Summernote
    soup = BeautifulSoup(document.content, 'html.parser')
    
    # Convert HTML to Word document elements
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'table', 'ul', 'ol']):
        if element.name == 'p':
            p = doc.add_paragraph()
            p.add_run(element.get_text())
        elif element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'table':
            # Convert HTML table to Word table
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text()
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{document.title}_{document.id}.docx"'
    return response


@staff_member_required
def certificate_list(request):
    """List all certificates for the Generate page"""
    documents = CertificateDocument.objects.all().order_by('-updated_at')
    return render(request, 'certificates/certificate_list.html', {'documents': documents})


@login_required
@require_http_methods(["GET", "POST"])
def edit_certificate_view(request, certificate_id):
    """Edit certificate content in the system before printing"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    # Check permission
    if certificate.created_by != request.user and not request.user.can_edit_records:
        messages.error(request, 'You do not have permission to edit this certificate.')
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = EditableCertificateForm(request.POST, instance=certificate)
        if form.is_valid():
            form.save()
            
            # Re-generate the Word document with edited content
            try:
                from .document_generator import CertificateDocumentGenerator
                generator = CertificateDocumentGenerator(certificate)
                generator.save_files(certificate)
            except ImportError:
                pass
            
            messages.success(request, 'Certificate updated successfully!')
            return redirect('certificates:preview', certificate_id=certificate.id)
    else:
        form = EditableCertificateForm(instance=certificate)
    
    # Convert certificate_data to HTML for editing if it's a dict
    if isinstance(certificate.certificate_data, dict):
        # Format the data as readable text
        html_content = '<div class="certificate-content">'
        for key, value in certificate.certificate_data.items():
            if value:
                html_content += f'<p><strong>{key.replace("_", " ").title()}:</strong> {value}</p>'
        html_content += '</div>'
        form.initial['certificate_data'] = html_content
    
    context = {
        'certificate': certificate,
        'form': form,
    }
    
    return render(request, 'certificates/edit_certificate.html', context)



@login_required
@require_http_methods(["GET"])
def print_certificate_from_system(request, certificate_id):
    """Display certificate as HTML for printing from the system"""
    certificate = get_object_or_404(GeneratedCertificate, id=certificate_id)
    
    # Check permission
    if certificate.created_by != request.user and not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to print this certificate.')
        return redirect('dashboard')
    
    # Render the HTML template with certificate data
    from django.template import Template, Context
    template = Template(certificate.template.html_template)
    context = Context(certificate.certificate_data)
    rendered_content = template.render(context)
    
    context = {
        'certificate': certificate,
        'data': certificate.certificate_data,
        'recipient_name': certificate.recipient_name,
        'certificate_id': certificate.id,
        'rendered_content': rendered_content,
    }
    
    return render(request, 'certificates/print_certificate.html', context)


# Requester views
@login_required
@require_http_methods(["GET"])
def request_certificate_types_view(request):
    """Display available certificate types for requesters"""
    if request.user.role != 'requester':
        messages.error(request, 'This page is for requesters only.')
        return redirect('dashboard')
    
    query = request.GET.get('q', '').strip()
    certificates = CertificateTemplate.objects.filter(is_active=True)
    
    if query:
        certificates = certificates.filter(
            Q(template_name__icontains=query) |
            Q(description__icontains=query) |
            Q(template_type__icontains=query)
        )
    
    # Define categories mapping
    category_mapping = {
        'barangay_clearance': 'Basic Residency Certificates',
        'residency': 'Basic Residency Certificates',
        'non_residency': 'Basic Residency Certificates',
        'indigency': 'Basic Residency Certificates',
        'Food_Assistance': 'Assistance Certificates',
        'Financial_Assistance': 'Assistance Certificates',
        'Program_Certificate': 'Program & Activity Certificates',
        'Participation': 'Program & Activity Certificates',
        'senior_citizen': 'Special Groups Certificates',
        'pwd': 'Special Groups Certificates',
        'solo_parent': 'Special Groups Certificates',
        'first_time_job_seeker': 'Special Groups Certificates',
        'death_certificate': 'Legal & Official Forms',
        'complaint_form': 'Legal & Official Forms',
        'hearing_notice': 'Legal & Official Forms',
        'settlement': 'Legal & Official Forms',
        'appearance': 'Legal & Official Forms',
        'duration': 'Legal & Official Forms',
        'business_permit': 'Business & Ownership Certificates',
        'animal_ownership': 'Business & Ownership Certificates',
        'certification': 'Health & Other Certificates',
        'other': 'Health & Other Certificates',
    }
    
    # Group certificates by category
    categorized_certificates = {}
    for cert in certificates:
        category = category_mapping.get(cert.template_type, 'Other Certificates')
        if category not in categorized_certificates:
            categorized_certificates[category] = []
        categorized_certificates[category].append(cert)
    
    context = {
        'certificates': certificates,
        'categorized_certificates': categorized_certificates,
        'query': query,
    }
    
    return render(request, 'certificates/request_types.html', context)


@login_required
@require_http_methods(["GET", "POST"])
def request_certificate_view(request, template_id):
    """Request a new certificate"""
    if request.user.role != 'requester':
        messages.error(request, 'This page is for requesters only.')
        return redirect('dashboard')
    
    template = get_object_or_404(CertificateTemplate, id=template_id, is_active=True)
    
    if request.method == 'POST':
        form = CertificateGenerationForm(request.POST)
        dynamic_form = DynamicCertificateForm(template, request.POST)
        
        if form.is_valid() and dynamic_form.is_valid():
            # Convert date objects to strings for JSON serialization
            request_data = {}
            for key, value in dynamic_form.cleaned_data.items():
                if isinstance(value, (date, datetime)):
                    request_data[key] = value.isoformat()
                else:
                    request_data[key] = value
            
            # Create certificate request
            cert_request = CertificateRequest.objects.create(
                template=template,
                requester=request.user,
                recipient_name=form.cleaned_data['recipient_name'],
                recipient_email=form.cleaned_data.get('recipient_email', ''),
                recipient_contact=form.cleaned_data.get('recipient_contact', ''),
                request_data=request_data,
                status='pending'
            )
            
            # Log the activity
            ActivityLog.objects.create(
                user=request.user,
                log_type='certificate_requested',
                action_description=f"Requested {template.template_name} for {cert_request.recipient_name}",
                object_id=cert_request.id,
            )
            
            messages.success(request, 'Certificate request submitted successfully!')
            return redirect('certificates:my_requests')
    else:
        form = CertificateGenerationForm()
        dynamic_form = DynamicCertificateForm(template)
    
    context = {
        'template': template,
        'form': form,
        'dynamic_form': dynamic_form,
    }
    
    return render(request, 'certificates/request_certificate.html', context)


@login_required
@require_http_methods(["GET"])
def my_requests_view(request):
    """Display current requester's certificate requests"""
    if request.user.role != 'requester':
        messages.error(request, 'This page is for requesters only.')
        return redirect('dashboard')

    query = request.GET.get('q', '').strip()
    requests = CertificateRequest.objects.filter(requester=request.user).select_related('template').order_by('-created_date')

    if query:
        requests = requests.filter(
            Q(template__template_name__icontains=query) |
            Q(recipient_name__icontains=query) |
            Q(status__icontains=query)
        )

    context = {
        'requests': requests,
        'query': query,
    }

    return render(request, 'certificates/my_requests.html', context)


@login_required
@require_http_methods(["GET"])
def manage_requests_view(request):
    """View and manage certificate requests (staff only)"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to manage requests.')
        return redirect('dashboard')
    
    requests = CertificateRequest.objects.select_related('template', 'requester').order_by('-created_date')
    
    # Filter by status if provided
    status = request.GET.get('status')
    if status:
        requests = requests.filter(status=status)
    
    context = {
        'requests': requests,
    }
    
    return render(request, 'certificates/manage_requests.html', context)


@login_required
@require_http_methods(["POST"])
def approve_request_view(request, request_id):
    """Approve a certificate request and generate certificate"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to approve requests.')
        return redirect('dashboard')
    
    cert_request = get_object_or_404(CertificateRequest, id=request_id)
    
    if cert_request.status != 'pending':
        messages.error(request, 'Request has already been processed.')
        return redirect('certificates:manage_requests')
    
    # Update request status
    cert_request.status = 'approved'
    cert_request.reviewed_by = request.user
    cert_request.reviewed_date = timezone.now()
    cert_request.save()
    
    # Generate certificate
    certificate = GeneratedCertificate.objects.create(
        template=cert_request.template,
        recipient_name=cert_request.recipient_name,
        recipient_email=cert_request.recipient_email,
        recipient_contact=cert_request.recipient_contact,
        certificate_data=cert_request.request_data,
        created_by=request.user,
        status='generated'
    )
    
    # Link to request
    cert_request.generated_certificate = certificate
    cert_request.status = 'completed'
    cert_request.save()
    
    # Generate files
    try:
        from .document_generator import CertificateDocumentGenerator
        generator = CertificateDocumentGenerator(certificate)
        generator.save_files(certificate)
    except ImportError:
        pass
    
    # Log activities
    ActivityLog.objects.create(
        user=request.user,
        log_type='certificate_generated',
        action_description=f"Generated certificate from request: {certificate.template.template_name} for {certificate.recipient_name}",
        object_id=certificate.id,
    )
    
    messages.success(request, 'Certificate request approved and certificate generated!')
    return redirect('certificates:manage_requests')


@login_required
@require_http_methods(["POST"])
def reject_request_view(request, request_id):
    """Reject a certificate request"""
    if not request.user.can_print_certificates:
        messages.error(request, 'You do not have permission to reject requests.')
        return redirect('dashboard')
    
    cert_request = get_object_or_404(CertificateRequest, id=request_id)
    
    if cert_request.status != 'pending':
        messages.error(request, 'Request has already been processed.')
        return redirect('certificates:manage_requests')
    
    reason = request.POST.get('rejection_reason', '').strip()
    
    # Update request status
    cert_request.status = 'rejected'
    cert_request.reviewed_by = request.user
    cert_request.reviewed_date = timezone.now()
    cert_request.rejection_reason = reason
    cert_request.save()
    
    messages.success(request, 'Certificate request rejected.')
    return redirect('certificates:manage_requests')