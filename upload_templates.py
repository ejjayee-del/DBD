#!/usr/bin/env python
"""
Helper script to upload certificate templates to the system.

Usage:
    python upload_templates.py --source /path/to/templates --destination media/templates
    python upload_templates.py --register      # Auto-register all templates
"""

import os
import sys
import shutil
from pathlib import Path
from docx import Document
import django

# Setup Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'dbd_project.settings')
django.setup()

from certificates.models import CertificateTemplate
import argparse


def extract_placeholders(docx_path):
    """Extract {{placeholder}} patterns from a docx file"""
    import re
    placeholders = set()
    pattern = r'\{\{([a-zA-Z0-9_]+)\}\}'
    
    try:
        doc = Document(docx_path)
        
        # Check paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            placeholders.update(matches)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        placeholders.update(matches)
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
    
    return sorted(list(placeholders))


def copy_templates(source_dir, dest_dir):
    """Copy .docx files from source to destination"""
    source_path = Path(source_dir)
    dest_path = Path(dest_dir)
    
    if not source_path.exists():
        print(f"❌ Source directory does not exist: {source_path}")
        return False
    
    dest_path.mkdir(parents=True, exist_ok=True)
    
    docx_files = list(source_path.glob('*.docx'))
    if not docx_files:
        print(f"❌ No .docx files found in {source_path}")
        return False
    
    for src_file in docx_files:
        dest_file = dest_path / src_file.name
        shutil.copy2(src_file, dest_file)
        print(f"✓ Copied: {src_file.name}")
    
    print(f"\n✓ Successfully copied {len(docx_files)} template(s) to {dest_path}")
    return True


def register_templates():
    """Register all templates from media/templates directory"""
    template_dir = Path('media/templates')
    
    if not template_dir.exists():
        print(f"❌ Template directory does not exist: {template_dir}")
        return
    
    docx_files = list(template_dir.glob('*.docx'))
    if not docx_files:
        print(f"❌ No .docx files found in {template_dir}")
        return
    
    print(f"\nRegistering {len(docx_files)} template(s)...")
    
    for filepath in docx_files:
        template_name = filepath.stem
        
        # Skip if already registered
        if CertificateTemplate.objects.filter(template_name=template_name).exists():
            print(f"⊙ Already registered: {template_name}")
            continue
        
        # Extract placeholders
        placeholders = extract_placeholders(filepath)
        print(f"{'✓' if placeholders else '!'} {template_name}")
        if placeholders:
            print(f"    Fields: {', '.join(placeholders)}")
        
        # Create template
        template = CertificateTemplate(
            template_type='other',
            template_name=template_name,
            template_file=f'templates/{filepath.name}',
            html_template=f'<!-- {template_name} -->',
            description=f'Auto-registered certificate template',
            is_active=True,
        )
        
        # Map required_fields from placeholders
        template.required_fields = {p: {'type': 'text', 'label': p.replace('_', ' ').title()} 
                                     for p in placeholders}
        template.save()
        
        # CertificateField records no longer needed - using JSON fields
        
        print(f"    ✓ Registered with {len(placeholders)} field(s)")
    
    print("\n✓ Template registration complete!")


def main():
    parser = argparse.ArgumentParser(description='Upload and register certificate templates')
    parser.add_argument('--source', help='Source directory with template files')
    parser.add_argument('--destination', default='media/templates', help='Destination directory')
    parser.add_argument('--register', action='store_true', help='Register all templates')
    
    args = parser.parse_args()
    
    if args.source:
        print(f"Copying templates from {args.source}...\n")
        if copy_templates(args.source, args.destination):
            print("\nWould you like to register these templates? (y/n): ", end='')
            if input().lower().startswith('y'):
                register_templates()
    elif args.register:
        register_templates()
    else:
        parser.print_help()


if __name__ == '__main__':
    main()
